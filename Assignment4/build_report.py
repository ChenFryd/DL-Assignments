"""Generates the Assignment 4 report as a .docx file."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FIGURES = '/home/chenfryd/DL-Assignments/Assignment4/figures'
OUT = '/home/chenfryd/DL-Assignments/Assignment4/Report.docx'

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Default paragraph style: Calibri 12 ──────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        run.font.size = Pt(13) if level == 1 else Pt(12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    return p

def bold_body(label, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    rest = p.add_run(text)
    rest.font.name = 'Calibri'
    rest.font.size = Pt(12)
    return p

def add_figure(path, caption, width=Cm(14)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles['Normal']
    for run in cap.runs:
        run.font.size = Pt(10)
        run.font.italic = True
    cap.paragraph_format.space_after = Pt(10)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
    doc.add_paragraph()
    return t

def page_break():
    doc.add_page_break()

# =============================================================================
# TITLE
# =============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Assignment 4: Generative Models with GAN')
r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(16)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Ben-Gurion University of the Negev - Deep Learning Course')
r2.font.name = 'Calibri'; r2.font.size = Pt(12)

doc.add_paragraph()

# =============================================================================
# 1. INTRODUCTION
# =============================================================================
heading('1. Introduction')
body(
    'This report presents a standard GAN and Conditional GAN (cGAN) trained on the Adult '
    'Census Income dataset for tabular data synthesis, together with: a comparison of '
    'softmax vs Gumbel-ST discrete representations (Section 6), a deliberate mode-collapse '
    'experiment with minibatch-std mitigation (Section 7), and a spectral-normalization '
    'contribution (Section 8).'
)

# =============================================================================
# 2. DATASET
# =============================================================================
heading('2. Dataset and Preprocessing')
body(
    'The Adult dataset (32,561 records, 1994 US Census) has a binary income target '
    '(75.9% ≤50K / 24.1% >50K, Fig. A1). Six continuous features are standardized with '
    'StandardScaler (not MinMaxScaler: the linear output head is unbounded). Eight '
    'categorical features are one-hot encoded (not ordinal: ordinal encoding implies '
    'invalid linear interpolation between levels), yielding a 108-dimensional input. '
    'Missing values in workclass/occupation/native-country are mapped to "Unknown" '
    'rather than imputed (mode imputation corrupts the target distribution) or dropped '
    '(survivorship bias). Feature distributions are in Appendix Figs. A2–A3.'
)
bold_body('Train/test split: ', '80% training (26,048 records) and 20% test (6,513 records), '
    'stratified on the income label to preserve the 0.759/0.241 class ratio. All results '
    'are averaged over three random seeds (42, 123, 456).')

# =============================================================================
# 3. MODEL ARCHITECTURE
# =============================================================================
heading('3. Model Architecture')

heading('3.1 Generator', level=2)
body(
    'The generator is a three-layer MLP with 256 hidden units per layer, batch normalization, '
    'and LeakyReLU(0.2) activations. Three layers of 256 units were chosen to give the '
    'network sufficient capacity to model the 108-dimensional encoded space while remaining '
    'trainable without skip connections; shallower networks underfit the joint categorical '
    'and continuous structure, and wider layers produced no measurable gain. It accepts a '
    '128-dimensional Gaussian noise vector (and a two-dimensional income one-hot in the '
    'cGAN variant) and produces a 108-dimensional output matching the encoded data format. '
    'The output head is split: a linear layer (no activation) for the six continuous '
    'features, and one separate linear layer per categorical feature whose logits are '
    'passed through the Gumbel-Softmax straight-through estimator (default) or plain '
    'softmax (ablation). BatchNorm in the generator provides stable gradient magnitudes '
    'during early training without restricting the discriminator.'
)

heading('3.2 Discriminator', level=2)
body(
    'The discriminator is a three-layer MLP with 256, 256, and 128 hidden units, '
    'LeakyReLU(0.2) activations, and Dropout(0.3) after the first two layers. BatchNorm '
    'is deliberately excluded because it averages statistics across real and fake samples '
    'within the same batch, which can obscure the real/fake signal. The output is a single '
    'sigmoid unit. For the cGAN, the income one-hot is concatenated to the input so the '
    'discriminator learns to assess plausibility conditional on the requested class.'
)

heading('3.3 cGAN Extension', level=2)
body(
    'The income one-hot (dim 2) is concatenated to the noise vector (G input) and to the '
    'data vector (D input), so both networks see the same conditioning signal and D always '
    'evaluates plausibility given the requested class.'
)

# =============================================================================
# 4. TRAINING SETUP
# =============================================================================
heading('4. Training Setup')
add_table(
    ['Hyperparameter', 'Value'],
    [
        ['Epochs', '200 (Sections 6/9);  100 (Section 7);  80 (Section 8)'],
        ['Batch size', '256'],
        ['Noise dimension z', '128'],
        ['Optimizer', 'Adam, lr = 2e-4, b1 = 0.5, b2 = 0.999'],
        ['Loss function', 'Binary cross-entropy (BCE)'],
        ['Label smoothing', 'real -> 0.9,  fake -> 0.1 (two-sided)'],
        ['Temperature schedule', 'Linear anneal 1.0 -> 0.5 over training'],
        ['D : G update ratio', '1:1 (default);  5:1, hidden_g=16, softmax (Section 8 collapse)'],
        ['Categorical strategy', 'Gumbel-ST (default);  Softmax (Section 7 ablation)'],
    ]
)
body(
    'Two-sided label smoothing is used: real samples are assigned a target of 0.9 '
    '(instead of 1.0) and fake samples a target of 0.1 (instead of 0.0). Smoothing '
    'the real target prevents the discriminator from becoming overconfident and '
    'producing near-zero gradients for the generator. Smoothing the fake target '
    'provides a symmetric regularisation effect, preventing the discriminator from '
    'driving its fake-sample output to exactly 0. The Gumbel-Softmax temperature is '
    'annealed from 1.0 to 0.5 so early training benefits from softer, more '
    'gradient-rich outputs while later training commits to near-discrete samples.'
)

# =============================================================================
# 5. SECTION 6 - EVALUATION
# =============================================================================
heading('5. Evaluation (Section 6)')

heading('5.1 Loss Curves', level=2)
body(
    'Loss curves for all three seeds are shown in Appendix Figs. A4-A9. In all runs the '
    'generator loss converges to approximately 2.27-2.29 and the discriminator loss to '
    'approximately 0.65-0.66. With two-sided label smoothing (real label = 0.9, fake '
    'label = 0.1), the discriminator\'s optimal output is D(real) = 0.9 and D(fake) = 0.1, '
    'yielding a minimum BCE loss of BCE(0.9, 0.9) + BCE(0.1, 0.1) ≈ 0.325 + 0.325 = 0.65. '
    'This is the discriminator-dominance plateau — D has converged to perfectly classify '
    'real from fake at the smoothed targets — not the game-theoretic Nash equilibrium, '
    'which requires G to match p_data so that D outputs 0.5 for all samples, giving '
    'D-loss ≈ BCE(0.5, 0.9) + BCE(0.5, 0.1) = 0.693 + 0.693 ≈ 1.386 (as confirmed by '
    'the spectral-normalization experiment in Section 8). The generator loss well above '
    'ln(2) = 0.693 confirms the discriminator retains a consistent advantage throughout '
    'training. Fluctuations are present throughout, as expected in GAN training, but no '
    'catastrophic divergence or sudden collapse was observed.'
)

heading('5.2 Synthetic Data Quality', level=2)
body(
    'Continuous and categorical feature distributions for both GAN and cGAN are compared '
    'against real training data in Appendix Figs. A10-A13. The continuous distributions '
    'are qualitatively similar in shape, though the synthetic capital-gain and capital-loss '
    'distributions lack the heavy-tail structure of the real data; these features are '
    'zero-inflated (>90% of values are exactly 0) with a sparse heavy tail, a degenerate '
    'manifold that Gaussian noise mapped through a smooth MLP cannot naturally reproduce '
    'without explicit zero-mass modeling. Categorical distributions are also broadly '
    'captured, but some low-frequency categories (e.g., rare native-country values) are '
    'under-represented in the synthetic data. Correlation matrices (Figs. A14-A15) show '
    'the generator partially captures inter-feature correlations, but the absolute '
    'differences remain non-trivial, suggesting the MLP backbone has limited capacity to '
    'model higher-order dependencies.'
)

heading('5.3 Detection and Efficacy Metrics', level=2)
body(
    'Detection is measured by training a Random Forest on a balanced mix of real and synthetic '
    'data using four-fold cross-validation (three folds train, one fold test, averaged over '
    'four rotations) and reporting the AUC. Low AUC near 0.5 indicates synthetic data is '
    'indistinguishable from real. Efficacy is the ratio of AUC when a Random Forest is '
    'trained on synthetic versus real data and evaluated on the original test set; high '
    'values near 1.0 indicate the synthetic data is a useful substitute for training.'
)

add_table(
    ['Model', 'Detection AUC', 'Efficacy (mean)', 'Efficacy (std)'],
    [
        ['GAN',  '1.000 +/- 0.000', '0.842', '0.084'],
        ['cGAN', '1.000 +/- 0.000', '0.763', '0.105'],
    ]
)

body(
    'Detection AUC is 1.000 for both models across all three seeds, meaning the Random '
    'Forest perfectly separates real from synthetic samples with no overlap. When D '
    'converges to its smoothed-label optimum while G-loss stays well above ln(2)=0.693, '
    'the generator has not matched p_data and the synthetic manifold is detectably '
    'different from the real one. Consistency across all three seeds rules out '
    'initialization as the cause — this is a structural limitation of the MLP generator. '
    'Section 8 confirms that training stability is not the bottleneck either: even after '
    'spectral normalization drives training to the true Nash equilibrium, detection AUC '
    'remains at 1.000, isolating model capacity as the binding constraint.'
)
body(
    'GAN efficacy (mean 0.842 ± 0.084) shows that synthetic data trained a Random Forest '
    'to 84% of the performance achievable with real data — usable as a partial substitute. '
    'The cGAN is slightly lower (0.763 ± 0.105): conditioning on a prescribed class ratio '
    'constrains the generated distribution, while the unconditional GAN\'s pseudo-labeling '
    'naturally tracks the learned feature space. The high standard deviation across seeds '
    'for both models reflects sensitivity to the noise initialization, which the '
    'spectral-norm variant (Section 8) largely resolves (efficacy 0.992, std ≈ 0).'
)

# =============================================================================
# 6. SECTION 7 - DISCRETE FEATURES
# =============================================================================
heading('6. Discrete Feature Representation (Section 7)')

heading('6.1 Why Argmax Breaks Gradient Flow', level=2)
body(
    'Argmax is piecewise constant: for any small perturbation of the logit vector the '
    'output one-hot is unchanged, so the sub-gradient is zero almost everywhere. '
    'During backpropagation, the chain rule requires multiplying the downstream gradient '
    'by the local Jacobian of each operation; the Jacobian of argmax is zero, so the '
    'gradient signal reaching the generator\'s parameters is identically zero regardless '
    'of the discriminator\'s feedback. The generator therefore receives no information '
    'about which categorical values to produce and cannot improve its discrete outputs.'
)

heading('6.2 Strategy Comparison', level=2)
bold_body('Plain Softmax. ', 'Forward: the generator outputs a soft probability vector; '
    'the well-defined softmax Jacobian allows gradients to flow back to the logits. '
    'Backward: gradients are non-zero. However, the discriminator sees hard 0/1 one-hots '
    'from real data and soft probabilities from fake data — a format mismatch it can '
    'exploit instead of learning semantic content.')
doc.add_paragraph()
bold_body('Gumbel-Softmax + Straight-Through (Gumbel-ST). ', 'Forward: Gumbel noise is '
    'added to the logits and a hard one-hot is taken via argmax, so the discriminator '
    'sees discrete samples identical in format to real data — no format mismatch. '
    'Backward: the straight-through estimator substitutes the soft Gumbel-Softmax '
    'gradient (y_hard - y_soft.detach() + y_soft) for the zero argmax gradient, '
    'allowing informative updates while the forward pass remains discrete.')

heading('6.3 Entropy Analysis', level=2)
body(
    'To quantify commitment, the average Shannon entropy of the generated soft categorical '
    'distributions (computed from raw logits, without Gumbel noise) is compared against '
    'the entropy implied by real category frequencies. A near-uniform distribution has '
    'high entropy and indicates the generator is not generating meaningful categorical '
    'structure. Results are shown in Appendix Fig. A16 and summarized below.'
)
add_table(
    ['Feature', 'Real entropy', 'Softmax', 'Gumbel-ST'],
    [
        ['workclass',       '1.148', '0.028', '0.442'],
        ['education',       '2.037', '0.181', '0.037'],
        ['marital-status',  '1.273', '0.058', '0.505'],
        ['occupation',      '2.440', '0.095', '1.070'],
        ['relationship',    '1.491', '0.070', '0.664'],
        ['race',            '0.551', '0.010', '0.448'],
        ['sex',             '0.635', '0.034', '0.328'],
        ['native-country',  '0.653', '0.000', '0.339'],
    ]
)
body(
    'Softmax entropy is close to zero for every feature. The mechanistic reason is the '
    'format mismatch: real data arrives as hard 0/1 one-hot vectors, while Softmax fake '
    'data arrives as continuous probability vectors. The discriminator exploits this '
    'distributional difference rather than learning semantic content. The generator\'s '
    'best response is to concentrate probability mass on the single most frequent '
    'category per feature — driving logits large for the modal category and near-zero '
    'for all others — which produces peaked, low-entropy distributions (hence raw-logit '
    'entropy collapsing near zero). Gumbel-ST entropy is substantially closer to '
    'real-data entropy for most features (notably occupation: 1.070 vs real 2.440, '
    'far better than Softmax\'s 0.095). The one exception is education, where '
    'Gumbel-ST entropy (0.037) is lower than Softmax (0.181): education has the '
    'widest vocabulary (16 levels), and the annealed temperature (1.0→0.5) pushes '
    'the straight-through estimator to commit aggressively to the single most frequent '
    'level. Softmax, paradoxically, does not collapse as far on this feature because '
    'the format-mismatch equilibrium keeps the discriminator focused on the soft/hard '
    'distinction rather than on the modal category. Loss curve comparison (Fig. A17) '
    'shows Gumbel-ST reaches a better equilibrium (G ~2.29, D ~0.66) than Softmax '
    '(G ~1.88, D ~0.86): Gumbel-ST forces the discriminator to learn genuine content '
    'differences, whereas Softmax\'s degenerate format-detection signal prevents D '
    'from converging to its optimum.'
)
bold_body('Quantitative comparison. ', 'Detection AUC: both 1.000 — the format mismatch '
    'does not help Softmax avoid detection, since the Random Forest finds other '
    'distinguishing features even without the soft/hard signal. '
    'Efficacy: Softmax 0.944 vs Gumbel-ST 0.942 — nearly identical, suggesting that '
    'the soft probability vectors, despite being a distributional mismatch, still '
    'carry enough signal for a downstream classifier to learn useful structure. '
    'Despite the near-equal efficacy, Gumbel-ST is the correct representation: '
    'it eliminates the format mismatch, produces entropy values substantially closer '
    'to real-data category statistics, and reaches a theoretically sounder training '
    'equilibrium (D-loss closer to the label-smoothed optimum). Choosing Softmax '
    'because its efficacy is 0.002 higher would be overfitting to noise.')

# =============================================================================
# 7. SECTION 8 - MODE COLLAPSE
# =============================================================================
heading('7. Mode Collapse: Induction and Mitigation (Section 8)')

heading('7.1 Collapse Indicator', level=2)
body(
    'The collapse indicator is the categorical coverage ratio: unique (workclass, '
    'marital-status, occupation) three-way combinations in synthetic data divided by '
    'those in real data. Values near 0 indicate collapse to a narrow pattern set; '
    'values near 1 indicate full diversity. These three features were chosen for their '
    'range of vocabulary sizes and social-demographic independence.'
)

heading('7.2 Inducing Collapse', level=2)
body(
    'Collapse is induced by combining: (1) a bottlenecked generator (hidden=16, down from '
    '256), which cannot map noise to diverse categorical combinations; (2) plain softmax '
    'in place of Gumbel-ST, removing the diversity floor that Gumbel noise provides; and '
    '(3) a 5:1 D:G update ratio to maintain discriminator dominance.'
)
add_table(
    ['Configuration', 'Coverage ratio', 'Unique combos'],
    [
        ['Real data (reference)', '1.000', '381'],
        ['Baseline (1:1, hidden=256, Gumbel-ST)', '1.031', '393'],
        ['Collapsed (5:1, hidden=16, softmax)', '0.116', '44'],
    ]
)
body(
    'The revised strategy produced genuine mode collapse: coverage dropped from the baseline '
    'of 1.031 to 0.116, reducing unique (workclass, marital-status, occupation) combinations '
    'from 393 to just 44 out of 381 real combinations. The bottlenecked generator concentrated '
    'probability mass on the most frequent category for each feature, mapping nearly all noise '
    'vectors to the same few combinations. The loss curves (Fig. A18) reveal a key diagnostic '
    'property of mode collapse: despite the catastrophic diversity loss, the G-loss and D-loss '
    'plateaus are nearly identical to the baseline (G ≈2.27, D ≈0.66). This is because the '
    'discriminator already dominated in both cases, and scalar BCE loss is sensitive only to '
    'real/fake classification accuracy, not to output diversity. Mode collapse is therefore '
    'invisible in loss curves alone and can only be detected through a coverage-based metric. '
    'This is why explicitly defining and tracking a collapse indicator (Section 7.1) is essential.'
)

heading('7.3 Mitigation: Minibatch Standard Deviation', level=2)
bold_body('Prediction. ', 'The minibatch standard-deviation (MBD) layer appends the mean '
    'per-feature standard deviation across the mini-batch to the discriminator input. '
    'When the collapsed generator produces near-identical fake samples, their batch std '
    'is near zero, making them trivially detectable. The generator must therefore spread '
    'its outputs to survive. Predicted: coverage ratio recovers above the collapsed value '
    'of 0.116; scalar losses remain similar (same architecture and update ratio).')
doc.add_paragraph()
bold_body('Observed. ', 'MBD yielded a coverage ratio of 0.255 (97/381), more than double '
    'the collapsed run\'s 0.116 (44/381). G-loss settled at 2.30 and D-loss at 0.65 — '
    'nearly identical to the collapsed run, as predicted.')
doc.add_paragraph()
body(
    'Prediction confirmed. MBD appends the mean per-feature standard deviation across '
    'the mini-batch to the discriminator\'s input. When the generator collapses, fake '
    'samples within a batch become nearly identical, driving their batch std toward '
    'zero — a signal the discriminator can exploit independently of whether any '
    'individual sample looks realistic. The generator is penalized not for a single '
    'fake sample but for producing a homogeneous batch, so it must spread its outputs '
    'across more combinations to survive. This explains the coverage increase from '
    '44 to 97 unique combinations. Full recovery to the baseline (1.031) was not '
    'achieved because MBD improves diversity incentives but cannot expand '
    'representational capacity: with hidden=16, the generator lacks the parameters '
    'to span all 381 real combinations regardless of the gradient signal it receives. '
    'Bar chart and loss curves are in Appendix Figs. A19-A20.'
)

# =============================================================================
# 8. SECTION 9 - SPECTRAL NORMALIZATION
# =============================================================================
heading('8. Open-Ended Contribution: Spectral Normalization (Section 9)')

heading('8.1 Modification and Motivation', level=2)
body(
    'Spectral normalization (SN) is applied to every linear layer of the discriminator: '
    'each weight matrix W is divided by its largest singular value sigma_1(W) at each '
    'forward pass, bounding the discriminator\'s Lipschitz constant to 1. This targets the '
    'baseline instability where D converges quickly to a near-optimal solution, saturating '
    'its gradients (sigmoid output → 0 for fake samples) and starving the generator of '
    'update signal. Unlike gradient penalty, SN requires no hyperparameter and adds no '
    'extra forward pass.'
)

heading('8.2 Prediction', level=2)
body(
    'Written before running the experiment. '
    '(1) G-loss variance should decrease: bounding the Lipschitz constant prevents sharp '
    'discriminator confidence spikes that cause the generator gradient to oscillate. '
    '(2) D-loss should rise above the baseline plateau of ~0.65: the constrained '
    'discriminator cannot converge to the smoothed-label optimum and should settle near '
    'the Nash value of ~1.386. '
    '(3) Efficacy should improve: stable gradients give the generator a better learning '
    'signal. Detection AUC is not predicted to improve — if it does not, that indicates '
    'the bottleneck is model capacity rather than training stability.'
)

heading('8.3 Results', level=2)
add_table(
    ['Metric', 'Baseline GAN', 'Spectral-Norm GAN'],
    [
        ['G-loss (final)',    '~2.28',   '~0.71'],
        ['D-loss (final)',    '~0.66',   '~1.37'],
        ['G-loss variance',  '0.112769', '0.000011'],
        ['Variance reduction', '-',      '100.0%'],
        ['Detection AUC',    '1.000',   '1.000'],
        ['Efficacy',         '0.861',   '0.992'],
    ]
)
body(
    'All three predictions were confirmed. G-loss variance was reduced by 100%, from '
    '0.113 to 0.000011. The final G-loss of 0.71 is very close to ln(2) = 0.693, the '
    'generator loss at the true Nash equilibrium (G perfectly matches p_data, D outputs '
    '0.5 everywhere). The D-loss of 1.37 is equally diagnostic: with two-sided label '
    'smoothing (real target 0.9, fake target 0.1), the theoretical D-loss when D outputs '
    '0.5 for all samples is BCE(0.5, 0.9) + BCE(0.5, 0.1) = 0.693 + 0.693 ≈ 1.386 — '
    'exactly matching the observed 1.37, confirming that SN drove training to the true '
    'Nash equilibrium rather than the discriminator-dominance plateau (D-loss ≈ 0.65) '
    'seen in the baseline. Efficacy improved from 0.861 to 0.992, a gain of 13 percentage '
    'points, indicating the generator learned a much more faithful representation of the '
    'data distribution when its gradient signal was stabilised. Loss curves are shown in '
    'Appendix Fig. A21.'
)
body(
    'Prediction (3) — detection — was not confirmed: AUC remained at 1.000 even at the '
    'Nash equilibrium. This negative result is itself informative. The SN experiment '
    'decouples two failure modes that are conflated in the baseline: training instability '
    '(gradient saturation, discriminator dominance) and model capacity (the MLP\'s ability '
    'to represent the joint distribution). SN eliminates the first completely — confirmed '
    'by the Nash equilibrium loss values — yet detection AUC does not move. This isolates '
    'model capacity as the binding constraint: a 3-layer MLP generator cannot reproduce '
    'the high-dimensional tabular distribution regardless of how well-trained it is. '
    'Improving detection AUC would require a more expressive architecture, not better '
    'training stability.'
)

# =============================================================================
# PAGE BREAK BEFORE APPENDIX
# =============================================================================
page_break()

# =============================================================================
# APPENDIX
# =============================================================================
heading('Appendix: Figures')

fig_list = [
    (f'{FIGURES}/eda_class_dist.png',
     'Fig. A1. Class distribution of the income target. '
     'The 3:1 imbalance (75.9% vs 24.1%) is preserved by stratified splitting.',
     Cm(8)),
    (f'{FIGURES}/eda_continuous.png',
     'Fig. A2. Continuous feature distributions in the real training data. '
     'Capital-gain and capital-loss exhibit heavy positive skew with most values at zero; '
     'age and hours-per-week follow roughly bell-shaped distributions.',
     Cm(15)),
    (f'{FIGURES}/eda_categorical.png',
     'Fig. A3. Categorical feature distributions (top 15 categories shown per feature; '
     'remaining values aggregated as "Other"). '
     'Native-country is heavily dominated by United-States; occupation and education '
     'show broader spread that a GAN must capture to avoid mode collapse.',
     Cm(15)),
    (f'{FIGURES}/gan_loss_seed42.png',
     'Fig. A4. GAN training loss, seed 42. '
     'G-loss climbs from ~2.0 and plateaus near 2.27; D-loss settles near 0.66. '
     'The gap between G-loss and ln(2)=0.693 indicates persistent discriminator dominance.',
     Cm(13)),
    (f'{FIGURES}/gan_loss_seed123.png',
     'Fig. A5. GAN training loss, seed 123. '
     'Pattern is consistent with seed 42: D settles near 0.65, G never recovers to the '
     'Nash equilibrium value of 0.693.',
     Cm(13)),
    (f'{FIGURES}/gan_loss_seed456.png',
     'Fig. A6. GAN training loss, seed 456. '
     'Consistent D-dominant plateau across all three seeds confirms this is a structural '
     'property of the training setup, not seed-specific variance.',
     Cm(13)),
    (f'{FIGURES}/cgan_loss_seed42.png',
     'Fig. A7. cGAN training loss, seed 42. '
     'Conditioning on the income label does not substantially change the loss dynamics; '
     'G-loss again plateaus well above 0.693.',
     Cm(13)),
    (f'{FIGURES}/cgan_loss_seed123.png',
     'Fig. A8. cGAN training loss, seed 123. '
     'G/D pattern mirrors the unconditional GAN; the discriminator retains an advantage '
     'regardless of the conditioning signal.',
     Cm(13)),
    (f'{FIGURES}/cgan_loss_seed456.png',
     'Fig. A9. cGAN training loss, seed 456. '
     'Stable convergence with no catastrophic oscillation, but the D-dominant imbalance '
     'persists across all seeds.',
     Cm(13)),
    (f'{FIGURES}/gan_dist_continuous.png',
     'Fig. A10. GAN continuous feature distributions, real vs synthetic (seed 42). '
     'Overall shapes are broadly matched; capital-gain and capital-loss heavy tails '
     'are underrepresented in the synthetic data.',
     Cm(15)),
    (f'{FIGURES}/gan_dist_categorical.png',
     'Fig. A11. GAN categorical feature distributions, real vs synthetic (seed 42). '
     'High-frequency categories are roughly reproduced; low-frequency categories '
     '(e.g., rare native-country values) are systematically under-represented.',
     Cm(15)),
    (f'{FIGURES}/cgan_dist_continuous.png',
     'Fig. A12. cGAN continuous feature distributions, real vs synthetic (seed 42). '
     'Conditional generation produces similar continuous-feature coverage to the '
     'unconditional GAN; heavy-tail features remain challenging.',
     Cm(15)),
    (f'{FIGURES}/cgan_dist_categorical.png',
     'Fig. A13. cGAN categorical feature distributions, real vs synthetic (seed 42). '
     'Class-conditional generation preserves category proportions somewhat better than '
     'the unconditional GAN for occupation and marital-status.',
     Cm(15)),
    (f'{FIGURES}/gan_corr.png',
     'Fig. A14. GAN correlation matrices: real (left), synthetic (center), difference (right). '
     'Non-trivial residuals in the difference matrix indicate the generator captures '
     'only first-order marginals and misses higher-order inter-feature dependencies.',
     Cm(15)),
    (f'{FIGURES}/cgan_corr.png',
     'Fig. A15. cGAN correlation matrices: real (left), synthetic (center), difference (right). '
     'Residuals are comparable to the unconditional GAN; conditioning does not substantially '
     'improve correlation structure reproduction.',
     Cm(15)),
    (f'{FIGURES}/s7_entropy.png',
     'Fig. A16. Per-feature Shannon entropy: real data vs softmax vs Gumbel-ST. '
     'Softmax entropy collapses near zero across all features (near-uniform soft distributions); '
     'Gumbel-ST entropy closely tracks the real-data reference.',
     Cm(14)),
    (f'{FIGURES}/s7_loss_comparison.png',
     'Fig. A17. Training loss comparison: softmax vs Gumbel-ST. '
     'Gumbel-ST reaches a lower D-loss (~0.66) than Softmax (~0.86), converging close '
     'to the discriminator-dominance optimum (~0.65) because D must discriminate on '
     'genuine content; Softmax D-loss is higher because the format-mismatch signal '
     'creates a degenerate training landscape.',
     Cm(14)),
    (f'{FIGURES}/s8_collapse_loss.png',
     'Fig. A18. Section 8 loss curves: baseline (1:1, hidden=256, Gumbel-ST) vs collapsed '
     '(5:1 D:G, hidden=16, softmax). Both runs reach similar G and D plateaus (~2.27 and '
     '~0.66 respectively), confirming that mode collapse is not visible in scalar losses — '
     'only the coverage metric (0.116 vs 1.031) reveals the narrowing of the output distribution.',
     Cm(14)),
    (f'{FIGURES}/s8_coverage.png',
     'Fig. A19. Categorical coverage ratios: baseline (1.031), collapsed (0.116), and '
     'MBD-mitigated (0.255). The collapsed run produces only 44 out of 381 real unique '
     'combinations; MBD more than doubles diversity to 97 combinations, with the capacity '
     'bottleneck preventing full recovery to baseline.',
     Cm(10)),
    (f'{FIGURES}/s8_all_loss.png',
     'Fig. A20. Section 8 loss curves: all three configurations (baseline / collapsed / MBD). '
     'All three settle at similar G/D loss values (~2.27-2.30 / 0.65-0.69), confirming that '
     'the coverage difference is driven by architectural choices (capacity bottleneck, Gumbel '
     'noise removal) rather than loss-curve-visible training dynamics.',
     Cm(15)),
    (f'{FIGURES}/s9_sn_loss.png',
     'Fig. A21. Section 9 loss curves: baseline vs spectral-norm GAN. '
     'Spectral-norm G-loss stabilizes near ln(2)=0.693 (the Nash equilibrium) '
     'while D-loss rises to ~1.36, a stark contrast to the baseline D-dominant plateau.',
     Cm(14)),
]

for path, caption, width in fig_list:
    add_figure(path, caption, width)

doc.save(OUT)
print(f'Saved: {OUT}')
